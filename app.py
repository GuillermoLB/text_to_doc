
from flask import Flask, request, send_file
from docx import Document
import spacy
import os

# Cargar el modelo de spaCy
nlp = spacy.load('en_core_web_sm')

app = Flask(__name__)

@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    data = request.json
    text = data.get('text')

    # Procesar el texto con NLP
    doc_nlp = nlp(text)
    entities = [(ent.text, ent.label_) for ent in doc_nlp.ents]

    # Crear un documento de Word
    doc = Document()
    doc.add_heading('Document Title', 0)

    doc.add_paragraph("Texto Original:")
    doc.add_paragraph(text)

    doc.add_heading('Entidades Extraídas', level=1)
    for entity, label in entities:
        doc.add_paragraph(f"{entity}: {label}")

    # Guardar el documento
    file_path = 'generated_document.docx'
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
